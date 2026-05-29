"""Convert hackernoon-launch.md to a Word .docx file.

Handles: H1-H4, paragraphs, **bold**, *italic*, `inline code`,
fenced code blocks (```python ... ```), tables, bullet/numbered lists,
blockquotes, [links](url), ![images](url), and horizontal rules.

Images at https://raw.githubusercontent.com/... are downloaded and embedded.
"""
import io
import re
import sys
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


SRC = Path(r"C:\Users\alexb\OneDrive\Рабочий стол\pmb\docs\blog\hackernoon-launch.md")
DST = Path(r"C:\Users\alexb\OneDrive\Рабочий стол\pmb\docs\blog\hackernoon-launch.docx")

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"

CODE_BG = RGBColor(0x0D, 0x11, 0x17)        # very dark for code blocks
CODE_FG = RGBColor(0xC9, 0xD1, 0xD9)        # light gray
LINK_BLUE = RGBColor(0x05, 0x66, 0xD9)
DIM = RGBColor(0x6E, 0x77, 0x81)


def set_shading(element, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0566D9")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


_INLINE_RE = re.compile(
    r"(`[^`]+`|"                              # inline code
    r"\*\*[^*]+\*\*|"                         # bold
    r"(?<!\*)\*[^*]+\*(?!\*)|"                # italic
    r"!\[[^\]]*\]\([^)]+\)|"                  # image (handled separately)
    r"\[[^\]]+\]\([^)]+\))"                   # link
)


def add_inline(paragraph, text, base_font=BODY_FONT, base_size=Pt(11)):
    """Parse inline markdown in `text` and add as runs to `paragraph`."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.name = base_font
            run.font.size = base_size
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(10)
            rPr = run._element.get_or_add_rPr()
            set_shading(rPr, "F1F3F5")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.name = base_font
            run.font.size = base_size
        elif token.startswith("[") and "](" in token:
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mm:
                add_hyperlink(paragraph, mm.group(2), mm.group(1))
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.name = base_font
            run.font.size = base_size
        else:
            run = paragraph.add_run(token)
            run.font.name = base_font
            run.font.size = base_size
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = base_font
        run.font.size = base_size


REPO_ROOT = Path(r"C:\Users\alexb\OneDrive\Рабочий стол\pmb")


def fetch_image(url):
    # First try mapping GitHub raw URL → local file
    m = re.match(
        r"https://raw\.githubusercontent\.com/oleksiijko/pmb/main/(.+)", url
    )
    if m:
        local = REPO_ROOT / m.group(1)
        if local.exists():
            print(f"  local: {local}")
            return local.read_bytes()
    # Fallback: download
    print(f"  fetching: {url}")
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=15) as r:
            return r.read()
    except Exception as e:
        print(f"  failed: {e}")
        return None


def setup_styles(doc):
    # Body font
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = HEADING_FONT
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    pf = h1.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = HEADING_FONT
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    pf = h2.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = HEADING_FONT
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    pf = h3.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)


def add_code_block(doc, lines, lang=""):
    """Render a code block as a dark box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    # Set background on the paragraph
    pPr = p._p.get_or_add_pPr()
    set_shading(pPr, "0D1117")

    # Add a thin border
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "30363D")
        pBdr.append(b)
    pPr.append(pBdr)

    code_text = "\n".join(lines)
    run = p.add_run(code_text)
    run.font.name = CODE_FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = CODE_FG


def add_table_from_md(doc, header, rows):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"

    # header
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = BODY_FONT
        tcPr = cell._tc.get_or_add_tcPr()
        set_shading(tcPr, "E8EEF6")

    # body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            text = row[c_idx] if c_idx < len(row) else ""
            add_inline(p, text.strip(), base_size=Pt(10.5))


def main():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_styles(doc)

    i = 0
    n = len(md)
    while i < n:
        line = md[i]
        stripped = line.strip()

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "D0D7DE")
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        # Image-only line: ![alt](url)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m:
            alt, url = m.group(1), m.group(2)
            data = fetch_image(url)
            if data:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(io.BytesIO(data), width=Inches(6.0))
                except Exception as e:
                    print(f"  embed failed: {e}")
                    p.add_run(f"[image: {alt}]").italic = True
            else:
                p = doc.add_paragraph()
                p.add_run(f"[image: {alt}] ({url})").italic = True
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            j = i + 1
            code_lines = []
            while j < n and not md[j].strip().startswith("```"):
                code_lines.append(md[j])
                j += 1
            add_code_block(doc, code_lines, lang)
            i = j + 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and md[i].strip().startswith(">"):
                ql = md[i].strip().lstrip(">").strip()
                quote_lines.append(ql)
                i += 1
            quote_text = " ".join([q for q in quote_lines if q])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            lft = OxmlElement("w:left")
            lft.set(qn("w:val"), "single")
            lft.set(qn("w:sz"), "18")
            lft.set(qn("w:space"), "8")
            lft.set(qn("w:color"), "0566D9")
            pBdr.append(lft)
            pPr.append(pBdr)
            set_shading(pPr, "F6F8FA")
            add_inline(p, quote_text)
            continue

        # Tables
        if stripped.startswith("|") and "|" in stripped[1:]:
            # gather all consecutive | lines
            tbl_lines = []
            while i < n and md[i].strip().startswith("|"):
                tbl_lines.append(md[i].strip())
                i += 1
            if len(tbl_lines) >= 2:
                def split_row(row):
                    parts = [c.strip() for c in row.strip("|").split("|")]
                    return parts
                header = split_row(tbl_lines[0])
                # separator is row 1 (---|---|---)
                body_rows = [split_row(r) for r in tbl_lines[2:]]
                add_table_from_md(doc, header, body_rows)
                doc.add_paragraph()  # spacing
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            while i < n and re.match(r"^\s*[-*]\s+", md[i]):
                text = re.sub(r"^\s*[-*]\s+", "", md[i])
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, text)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < n and re.match(r"^\s*\d+\.\s+", md[i]):
                text = re.sub(r"^\s*\d+\.\s+", "", md[i])
                p = doc.add_paragraph(style="List Number")
                add_inline(p, text)
                i += 1
            continue

        # Blank
        if stripped == "":
            i += 1
            continue

        # Regular paragraph: gather continuation lines
        para_lines = [line]
        i += 1
        while i < n:
            nxt = md[i]
            nxt_strip = nxt.strip()
            if (nxt_strip == ""
                or nxt_strip.startswith("#")
                or nxt_strip.startswith("```")
                or nxt_strip.startswith(">")
                or nxt_strip.startswith("|")
                or nxt_strip == "---"
                or re.match(r"^[-*]\s+", nxt_strip)
                or re.match(r"^\d+\.\s+", nxt_strip)
                or re.match(r"^!\[", nxt_strip)):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, text)

    doc.save(DST)
    print(f"\nWrote: {DST}")
    print(f"Size: {DST.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
